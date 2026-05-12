"""Word (DOCX) redline generation and PDF executive summary using reportlab."""
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)


def generate_redline_docx(contract_id: str, redlines: list[dict]) -> bytes:
    """
    Build a Word document with original → proposed clause rewrites.
    redlines: list of {clause_type, original_text, proposed_text, rationale}
    Returns DOCX as bytes.
    """
    doc = Document()
    doc.add_heading("ContractIQ — Redline Report", 0)
    doc.add_paragraph(f"Contract ID: {contract_id}")
    doc.add_paragraph("")

    for item in redlines:
        doc.add_heading(item.get("clause_type", "Clause").replace("_", " ").title(), level=2)

        # Original (red strikethrough style)
        orig_para = doc.add_paragraph()
        orig_para.add_run("Original: ").bold = True
        run = orig_para.add_run(item.get("original_text", ""))
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        # Proposed (green)
        prop_para = doc.add_paragraph()
        prop_para.add_run("Proposed: ").bold = True
        run2 = prop_para.add_run(item.get("proposed_text", ""))
        run2.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

        # Rationale
        rat_para = doc.add_paragraph()
        rat_para.add_run("Rationale: ").bold = True
        rat_para.add_run(item.get("rationale", ""))

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf_summary(risk_report: dict, clause_bundle: dict) -> bytes:
    """
    Build a PDF executive summary with risk heatmap table.
    Returns PDF as bytes.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    # Title
    title_style = ParagraphStyle("title", parent=styles["Title"],
                                 fontSize=20, spaceAfter=6)
    story.append(Paragraph("ContractIQ — Executive Risk Summary", title_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1a56db")))
    story.append(Spacer(1, 0.4*cm))

    # Meta
    contract_id = risk_report.get("contract_id", "—")
    overall = risk_report.get("contract_risk_score", 0)
    action = risk_report.get("recommended_action", "—")
    story.append(Paragraph(f"<b>Contract ID:</b> {contract_id}", styles["Normal"]))
    story.append(Paragraph(f"<b>Overall Risk Score:</b> {overall}/100", styles["Normal"]))
    story.append(Paragraph(f"<b>Recommended Action:</b> {action}", styles["Normal"]))
    story.append(Spacer(1, 0.4*cm))

    # Executive summary
    story.append(Paragraph("<b>Executive Summary</b>", styles["Heading2"]))
    story.append(Paragraph(risk_report.get("executive_summary", ""), styles["Normal"]))
    story.append(Spacer(1, 0.4*cm))

    # Clause risk table
    story.append(Paragraph("<b>Clause Risk Breakdown</b>", styles["Heading2"]))
    table_data = [["Clause", "Score", "Level", "Action"]]
    for cr in risk_report.get("clause_risks", []):
        level = cr.get("risk_level", "LOW")
        color_map = {"HIGH": colors.HexColor("#fef2f2"),
                     "MEDIUM": colors.HexColor("#fffbeb"),
                     "LOW": colors.HexColor("#ecfdf5")}
        table_data.append([
            cr.get("clause_type", "").replace("_", " ").title(),
            str(cr.get("risk_score", 0)),
            level,
            cr.get("recommended_action", ""),
        ])

    t = Table(table_data, colWidths=[8*cm, 2*cm, 3*cm, 4*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a56db")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)

    doc.build(story)
    return buf.getvalue()
